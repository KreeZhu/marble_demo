from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "GAME_DESIGN_DOCUMENT.md"
OUTPUT = ROOT / "docs" / "弹珠_游戏设计文档_v1.0.docx"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft YaHei"
NAVY = "1F4D78"
BLUE = "2E74B5"
GREEN = "16845B"
PURPLE = "7357B8"
AMBER = "A86D12"
MUTED = "64748B"
TEXT = "172033"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E9F6F0"
LIGHT_PURPLE = "F1ECFA"
WHITE = "FFFFFF"


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def configure_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(index, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name=FONT_LATIN) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color=TEXT, bold=False) -> None:
    style.font.name = FONT_LATIN
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def add_numbering_definition(document: Document, marker: str) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if marker == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if marker == "bullet" else "%1.")
    level.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_inline(paragraph, text: str, size=11, color=TEXT) -> None:
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=max(9, size - 0.5), color=NAVY, name="Consolas")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def set_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def add_rule(paragraph, color=BLUE, size=8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_document() -> tuple[Document, int, int]:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    set_style_font(styles["Title"], 30, NAVY, True)
    styles["Title"].paragraph_format.space_after = Pt(8)
    set_style_font(styles["Subtitle"], 14, MUTED, False)
    styles["Subtitle"].paragraph_format.space_after = Pt(18)

    header = section.header.paragraphs[0]
    header.clear()
    left = header.add_run("《弹珠》游戏设计文档")
    set_run_font(left, size=9, color=NAVY, bold=True)
    right = header.add_run("    规则基线 · 2026-08-13")
    set_run_font(right, size=9, color=MUTED)
    add_rule(header, color=LIGHT_BLUE, size=6)
    set_page_field(section.footer.paragraphs[0])

    bullet_num_id = add_numbering_definition(doc, "bullet")
    decimal_num_id = add_numbering_definition(doc, "decimal")
    return doc, bullet_num_id, decimal_num_id


def add_cover(doc: Document) -> None:
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("GAME DESIGN DOCUMENT · V1.0")
    set_run_font(run, size=10, color=GREEN, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("弹 珠")
    set_run_font(title_run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("二维物理解谜游戏 · 玩法规则与制作基线")
    set_run_font(subtitle_run, size=14, color=MUTED)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.left_indent = Inches(1.15)
    rule.paragraph_format.right_indent = Inches(1.15)
    add_rule(rule, color=GREEN, size=12)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.left_indent = Inches(0.65)
    lead.paragraph_format.right_indent = Inches(0.65)
    lead.paragraph_format.space_before = Pt(18)
    lead.paragraph_format.space_after = Pt(28)
    add_inline(
        lead,
        "固定力度，自由调角度。用反弹、机关、传送和可接球发射器，\n让同一颗弹珠最终碰到 B 点。",
        size=13,
        color=TEXT,
    )

    metadata = [
        ("对应游戏版本", "v0.1.0 当前开发版"),
        ("规则基线日期", "2026-08-13"),
        ("适用对象", "策划、开发、美术、音效、测试与非开发协作者"),
        ("文档目的", "统一机制理解，并作为后续逐条确认与验收的基准"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Table Grid"
    configure_table_geometry(table, [2700, 6660])
    for row_index, (label, value) in enumerate(metadata):
        label_cell, value_cell = table.rows[row_index].cells
        set_cell_shading(label_cell, LIGHT_BLUE)
        label_p = label_cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        label_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        label_run = label_p.add_run(label)
        set_run_font(label_run, size=10, color=NAVY, bold=True)
        value_p = value_cell.paragraphs[0]
        value_p.paragraph_format.space_after = Pt(0)
        value_run = value_p.add_run(value)
        set_run_font(value_run, size=10, color=TEXT)

    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def table_widths(rows: list[list[str]]) -> list[int]:
    columns = len(rows[0])
    header = rows[0]
    if columns == 2:
        return [2700, 6660]
    if columns == 3:
        return [1900, 2500, 4960]
    if columns == 4 and header[0] == "关卡":
        return [720, 1500, 2580, 4560]
    if columns == 4:
        return [1600, 1600, 2300, 3860]
    base = CONTENT_WIDTH_DXA // columns
    result = [base] * columns
    result[-1] += CONTENT_WIDTH_DXA - sum(result)
    return result


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    widths = table_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    configure_table_geometry(table, widths)
    mark_repeat_header(table.rows[0])
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif rows[0][0] == "关卡" and row_index % 2 == 0:
                set_cell_shading(cell, "F7F9FC")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.08
            if col_index == 0 and len(rows[0]) >= 3:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(paragraph, value, size=9.2 if len(rows) > 12 else 9.8, color=NAVY if row_index == 0 else TEXT)
            for run in paragraph.runs:
                if row_index == 0:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F6F9")
    p_pr.append(shd)
    for index, line in enumerate(code_lines):
        run = paragraph.add_run(line)
        set_run_font(run, size=8.5, color=NAVY, name="Consolas")
        if index < len(code_lines) - 1:
            run.add_break(WD_BREAK.LINE)


def add_callout(doc: Document, text: str, kind="key") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    configure_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GREEN if kind == "key" else LIGHT_PURPLE)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    add_inline(paragraph, text, size=11, color=GREEN if kind == "key" else PURPLE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_body_from_markdown(doc: Document, bullet_id: int, decimal_id: int) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## 1."))
    lines = lines[start:]
    index = 0
    in_code = False
    code_lines: list[str] = []
    active_decimal_id = decimal_id
    in_number_list = False

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if line.startswith("```"):
            in_number_list = False
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not line or line == "---":
            index += 1
            continue
        if line.startswith("|"):
            in_number_list = False
            rows, next_index = parse_table(lines, index)
            add_table(doc, rows)
            index = next_index
            continue
        if line.startswith("## "):
            in_number_list = False
            paragraph = doc.add_paragraph(style="Heading 1")
            paragraph.add_run(line[3:])
            for run in paragraph.runs:
                set_run_font(run, size=16, color=BLUE, bold=True)
            index += 1
            continue
        if line.startswith("### "):
            in_number_list = False
            paragraph = doc.add_paragraph(style="Heading 2")
            paragraph.add_run(line[4:])
            for run in paragraph.runs:
                set_run_font(run, size=13, color=BLUE, bold=True)
            index += 1
            continue
        bullet_match = re.match(r"^-\s+(.*)$", line)
        number_match = re.match(r"^\d+\.\s+(.*)$", line)
        if bullet_match:
            in_number_list = False
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            apply_numbering(paragraph, bullet_id)
            bullet_text = bullet_match.group(1)
            if bullet_text.startswith("[ ] "):
                bullet_text = "☐ " + bullet_text[4:]
            add_inline(paragraph, bullet_text)
            index += 1
            continue
        if number_match:
            if not in_number_list:
                active_decimal_id = add_numbering_definition(doc, "decimal")
                in_number_list = True
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            apply_numbering(paragraph, active_decimal_id)
            add_inline(paragraph, number_match.group(1))
            index += 1
            continue

        if line in {"**没有任何发射器会在接球后自动发射。**", "**胜利判定优先于其他接触结果。**"}:
            in_number_list = False
            add_callout(doc, line.strip("*"), "key")
            index += 1
            continue

        in_number_list = False
        paragraph = doc.add_paragraph()
        add_inline(paragraph, line.replace("  ", ""))
        index += 1


def audit_document(doc: Document) -> None:
    section = doc.sections[0]
    assert int(section.page_width.twips) == PAGE_WIDTH_DXA
    assert int(section.left_margin.twips) == 1440
    assert int(section.right_margin.twips) == 1440
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == CONTENT_WIDTH_DXA
        assert tbl_ind is not None and int(tbl_ind.get(qn("w:w"))) == TABLE_INDENT_DXA
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        assert not text.startswith("- "), f"Fake bullet found: {text}"


def main() -> None:
    doc, bullet_id, decimal_id = setup_document()
    doc.core_properties.title = "《弹珠》游戏设计文档"
    doc.core_properties.subject = "二维物理解谜游戏玩法规则与制作基线"
    doc.core_properties.author = "弹珠项目组"
    doc.core_properties.keywords = "弹珠, 游戏设计文档, 物理解谜, 关卡编辑器"
    add_cover(doc)
    add_body_from_markdown(doc, bullet_id, decimal_id)
    audit_document(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
